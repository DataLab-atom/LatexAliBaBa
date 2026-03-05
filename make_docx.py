"""
生成阿里巴巴 AI 研究资助计划 2026 申请书 Word 版
改进版：对齐 PDF 格式，包含页眉页脚、1.15 行距、表格优化等
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面设置 ──────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width  = Cm(21.0)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)
section.different_first_page_header_footer = True  # 封面无页眉

# ── 默认正文样式（1.15 行距 + 6pt 段后间距） ────────────────────────────────
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def _apply_line_spacing(p, multiple=1.15, space_after_pt=6):
    """对段落设置 1.15 行距和段后间距"""
    pPr = p._p.get_or_add_pPr()
    # 删除旧的 spacing 元素（如果有）
    for old in pPr.findall(qn('w:spacing')):
        pPr.remove(old)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(multiple * 240)))   # 1.15*240=276
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:after'), str(int(space_after_pt * 20)))   # pt -> twips
    pPr.append(spacing)


# ── 页眉（首页除外） ──────────────────────────────────────────────────────────
header = section.header
header.is_linked_to_previous = False
if header.paragraphs:
    hp = header.paragraphs[0]
else:
    hp = header.add_paragraph()
hp.clear()
hp.paragraph_format.tab_stops.add_tab_stop(Cm(14.4), WD_ALIGN_PARAGRAPH.RIGHT)
rl = hp.add_run('阿里巴巴 AI 研究资助计划 2026')
rl.font.size = Pt(9)
rl.font.name = '宋体'
rl._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
hp.add_run('\t')
rr = hp.add_run('课题一 & 十')
rr.font.size = Pt(9)
rr.font.name = '宋体'
rr._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# 页眉下边框
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '000000')
pBdr.append(bot)
hp._p.get_or_add_pPr().append(pBdr)

# ── 页脚（居中页码） ──────────────────────────────────────────────────────────
footer = section.footer
footer.is_linked_to_previous = False
if footer.paragraphs:
    fp = footer.paragraphs[0]
else:
    fp = footer.add_paragraph()
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = fp.add_run()
run_f.font.size = Pt(9)
for tag, text in [('begin', ''), ('', 'PAGE'), ('end', '')]:
    if tag in ('begin', 'end'):
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), tag)
        run_f._r.append(fc)
    else:
        it = OxmlElement('w:instrText')
        it.text = text
        run_f._r.append(it)

# ── 辅助：设置字体 ───────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, name='宋体'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)


# ── 辅助：表格单元格底纹 ──────────────────────────────────────────────────────
def shade_cell(cell, color='E8E8E8'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


# ── 辅助：表格单元格内边距 ────────────────────────────────────────────────────
def set_cell_margins(cell, top=60, start=100, bottom=60, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:w'), str(val))
        elem.set(qn('w:type'), 'dxa')
        tcMar.append(elem)
    tcPr.append(tcMar)


# ── 辅助：设置表格总宽度 ─────────────────────────────────────────────────────
def set_table_width(table, width_cm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_cm * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


# ── 辅助：对表格所有单元格应用内边距和行距 ───────────────────────────────────
def format_table(table, header_shade='D4D4D4', body_shade=None):
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            if i == 0:
                shade_cell(cell, header_shade)
            elif body_shade:
                shade_cell(cell, body_shade)
            for para in cell.paragraphs:
                _apply_line_spacing(para, multiple=1.15, space_after_pt=0)


# ── 段落样式函数 ─────────────────────────────────────────────────────────────

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    _apply_line_spacing(p, 1.15, 6)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 下划线分隔（与 PDF hrule 对应）
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    _apply_line_spacing(p, 1.15, 4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p


def heading3(text):
    """对应 LaTeX subsubsection*"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    _apply_line_spacing(p, 1.15, 3)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '黑体'
    run.font.italic = False
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p


def body(text, indent=False, bold_prefix=None):
    """正文段落，可选首段加粗前缀"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    _apply_line_spacing(p, 1.15, 6)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.bold = True
        rb.font.size = Pt(11)
        rb.font.name = '宋体'
        rb._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    _apply_line_spacing(p, 1.15, 3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def numbered(text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _apply_line_spacing(p, 1.15, 4)
    run = p.add_run(f'{num}.\u2003{text}')
    run.font.size = Pt(11)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


def code_block(text):
    """仿 LaTeX alignat 公式块，带框线背景"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    _apply_line_spacing(p, 1.15, 6)
    # 给段落加方框（四周细线）
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:space'), '4')
        elem.set(qn('w:color'), '888888')
        pBdr.append(elem)
    pPr.append(pBdr)
    # 浅灰底色
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p


# ══════════════════════════════════════════════════════════════════════════════
#  封面（首页，无页眉页脚）
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after  = Pt(12)
_apply_line_spacing(p, 1.15, 12)
run = p.add_run('阿里巴巴 AI 研究资助计划 2026')
run.font.size = Pt(20)
run.font.bold = True
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(36)
_apply_line_spacing(p2, 1.15, 36)
run2 = p2.add_run('研究申请书')
run2.font.size = Pt(16)
run2.font.bold = True
run2.font.name = '黑体'
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 封面信息表（居中，带框线）
cover_data = [
    ('项目名称', '基于大模型冷启动蒸馏的专家小智能体自动化生产\n——面向 AIOps 大小模型协同故障应急'),
    ('课题方向', '课题一（大小模型协同）/ 课题十（AIOps）'),
    ('日期',     '2026年3月'),
    ('申请机构', '香港城市大学  计算机科学系'),
    ('申请人（负责人）', '王茂林（Maolin Wang）'),
    ('电子邮件', 'maolwang2@cityu.edu.hk'),
]
tbl = doc.add_table(rows=len(cover_data), cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'
set_table_width(tbl, 14.0)
col_widths = [Cm(4.5), Cm(9.5)]
for i, (k, v) in enumerate(cover_data):
    row = tbl.rows[i]
    row.cells[0].width = col_widths[0]
    row.cells[1].width = col_widths[1]
    for cell in row.cells:
        set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
    shade_cell(row.cells[0], 'F0F0F0')
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(k)
    r0.font.bold = True; r0.font.size = Pt(11); r0.font.name = '宋体'
    r0._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    _apply_line_spacing(p0, 1.15, 0)
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(v)
    r1.font.size = Pt(11); r1.font.name = '宋体'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    _apply_line_spacing(p1, 1.15, 0)

# 大间距
for _ in range(4):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
_apply_line_spacing(p_note, 1.15, 0)
rn = p_note.add_run('本申请书响应 Alibaba AI Research Scheme 2026 资助计划。\n资助金额：人民币 60 万元，项目周期：一年。')
rn.font.size = Pt(10)
rn.font.italic = True
rn.font.name = '宋体'
rn._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. 项目介绍
# ══════════════════════════════════════════════════════════════════════════════
heading1('1.  项目介绍')

heading2('1.1  项目标题')
p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 6)
r = p.add_run('基于大模型冷启动蒸馏的专家小智能体自动化生产——面向 AIOps 大小模型协同故障应急')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

heading2('1.2  背景与意义')

body(
    '阿里巴巴等大规模互联网平台须保障近乎零停机时间。以电商大促为例，单次高峰期可触发数万条告警，'
    '人工 triage 每条平均耗时 5–15 分钟。设故障事件 ℰ = (x, a*) 为优化对象，'
    '课题一（大小模型协同）与课题十（AIOps）在最小化端到端响应延迟的目标中天然交汇：'
    '大模型 Planner（τ_L ∈ [500ms, 2s]）负责分解与路由，专家小模型（τ_S < 100ms）负责窄域深度执行。',
    bold_prefix='【问题建模：AIOps 故障应急的大小模型协同优化】\n'
)

body(
    '现有协同框架（FrugalGPT、RouteLLM）仅优化路由策略，均依赖"所有专家小模型已存在"这一前提，'
    '却未回答：对新出现的 AIOps 子场景，专家小模型从何而来？\n'
    '现有 SLM 生产方法——基于人工标注 CoT 的监督微调、知识蒸馏或自举过滤——均依赖以下至少一项：'
    '(a) 人工标注 CoT；(b) 具备领域知识的教师模型；(c) 已有非平凡解题能力的学生模型。'
    '而在冷启动场景（新 AIOps 子场景的常态），上述资源均不可得。',
    bold_prefix='【冷启动困境：专家小模型从何而来？】\n'
)

body(
    '如何在零人工标注、无领域感知教师的严格冷启动条件下，仅从原始 ⟨输入, 输出⟩ 对自动生产专家 SLM，'
    '使大小模型协同优化的联合目标得以完整求解？\n'
    '规则引擎无法泛化至新故障类型，ML 分类器缺乏可解释推理步骤，'
    'RAG 每次仍须调用云端 LLM 而无法消除延迟瓶颈——均未解决这一研究空白。',
    bold_prefix='【核心问题】\n'
)

body(
    '三点使之成为最佳选择：\n'
    '① 运维数据库中积累的 ⟨告警+日志, 修复操作⟩ 对天然满足所需格式，无需额外标注；\n'
    '② 修复有效性既可通过历史记录字符串匹配验证，也可在故障回放沙箱中以执行结果直接衡量；\n'
    '③ 本地 SLM 的 <100ms 延迟与云端 API 的 500ms–2s 形成直接量化对比，指标清晰可复现。',
    bold_prefix='【AIOps 是理想验证场景】\n'
)

heading2('1.3  研究目标')
body('本项目设定四项紧密关联的研究目标：')
numbered(
    '【O1】冷启动专家 SLM 生产框架。\n'
    '设计并验证一套三循环迭代流水线，仅从原始 ⟨输入, 输出⟩ 对出发，经由答案验证 CoT 合成（Cycle 1）、'
    '监督推理训练（Cycle 2）、领域聚焦策略优化（Cycle 3）及自举（Bootstrap）回路，'
    '实现冷启动条件下专家 SLM 的全自动化生产。', 1)
numbered(
    '【O2】AIOps 场景系统验证。\n'
    '在阿里巴巴运维数据集及公开 AIOps 基准上验证所提流水线，'
    '证明本地专家 SLM 在故障诊断与修复准确率上能达到或接近通用大模型的水准。', 2)
numbered(
    '【O3】专业化机制实证分析。\n'
    '通过消融实验（变量为模型规模、CoT 过滤率）及表示空间分析，'
    '对"优质 CoT 训练使小模型在窄域任务上超越大模型"现象进行实证论证。', 3)
numbered(
    '【O4】跨域适用性初探（探索性）。\n'
    '在资源允许时，将流水线应用于 AIOps 以外的公开数据集，初步验证本方法的领域无关性；'
    '否则以消融实验形式作为后续工作铺垫。', 4)

body(
    'SLM 修复准确率与 GPT-4 级 LLM 差距 ≤5%，推理延迟降低 >10×（<100ms），'
    'token 成本降至零（本地推理）；投稿 CCF-A 论文 2 篇（争取 1 篇结题前录用）；'
    '申请发明专利 1–2 项；2 个以上 AIOps 子场景完成原型部署与性能验证。',
    bold_prefix='【量化目标】'
)

heading2('1.4  研究方法')

heading3('核心思路：三循环迭代答案验证 CoT 合成')

body(
    '本项目拟设计一套三循环迭代 CoT 合成流水线，领域无关，'
    '无需任何人工标注或具备领域知识的教师模型，将原始 ⟨输入, 输出⟩ 对转化为具有推理能力的专家 SLM。'
    '系统由三个相互依存的训练循环构成，并配有一个自举（Bootstrap）回路：'
)

# 三循环公式块（对应 LaTeX alignat 环境）
code_block(
    'Cycle 1:  ⟨q, a⟩（冷启动数据）\n'
    '              →[最小教师 + k 采样 + 答案过滤]→  D*（验证三元组集）\n\n'
    'Cycle 2:  D*\n'
    '              →[SFT + LoRA]→  π₀（初始 SLM）\n\n'
    'Cycle 3:  π₀\n'
    '              →[AIOps-GRPO]→  π*（专家 SLM）\n\n'
    'Bootstrap: πₜ(D∅) →[过滤]→ D*_new ↪ D*\n'
    '              →[触发下一轮 Cycle 1]→ πₜ₊₁'
)

body(
    '关键在于：以历史记录中已有的标准答案作为唯一外部验证信号，无需引入额外人工监督，'
    '即可自动筛选通向正确答案的推理链，从而将通用模型的生成能力转化为经验证的领域特定训练数据。'
)

heading3('三循环训练流程')

body(
    '从运维日志中收集 ⟨告警+日志, 修复操作⟩ 对，无需任何人工标注。\n'
    '针对每条查询，由任务可行的最小外部教师模型（温度采样 τ=0.8，top-p=0.95）生成 k 条候选推理链；'
    '选用最小可用模型可将生成成本降低约 5×–10×。\n'
    '对每条候选推理链，提取末尾预测操作序列与标准答案进行集合等价匹配（规范化后），'
    '构建验证三元组集 D* = {(q, cⱼ, a) | ans(cⱼ) ≡ a}。'
    '保留率通常在 30–70% 区间（k=8）。',
    bold_prefix='Cycle 1——数据采集与验证 CoT 合成\n'
)

body(
    '将 D* 中的验证三元组用于对紧凑型 SLM（简单子任务选用 1.5B，复杂子任务选用 4B）'
    '进行带 LoRA 的监督微调，最大化 CoT 与答案的联合自回归似然。'
    '损失函数对推理链过短的样本降权，防止模型旁路推理步骤直接输出答案。'
    '此阶段产出初始 SLM π₀。',
    bold_prefix='Cycle 2——有监督推理训练（SFT + LoRA）\n'
)

body(
    '对 π₀ 施加群体相对策略优化：每条 prompt 生成 8 条 rollout，'
    '按奖励计算组内相对优势并更新策略参数。\n\n'
    'AIOps 奖励分为两层：\n'
    '  ▸ 轻量符号层（即时计算）：格式奖励 R_format、推理链长度奖励 R_CoT、精确匹配奖励 R_EM、'
    '软匹配奖励 R_soft、长度惩罚 R_len（负项）。\n'
    '  ▸ 执行结果层（沙箱异步回放）：内存峰值恢复率 R_mem、服务心跳通过率 R_svc、'
    '错误率下降比 R_err、p99 延迟恢复比 R_lat。\n\n'
    '总奖励 R_total = λ_s × R_sym + λ_e × R_exec；训练早期以符号层为主，中后期逐步增大 λ_e；'
    'KL 散度惩罚限制策略偏离初始点以保持训练稳定。',
    bold_prefix='Cycle 3——领域聚焦策略优化（AIOps-GRPO）\n'
)

body(
    '当某子任务在 Cycle 1 过滤率为零（D* = ∅）时，用当前已训练的 SLM（πₜ）对失败样本重新采样，'
    '经答案验证过滤后写回 D*，触发下一轮迭代，实现数据集与模型能力的协同进化。'
    '在积累足够 Bootstrap 样本前，该子任务请求暂由 Planner LLM 直接处理，维持系统可用性。',
    bold_prefix='Bootstrap 回路——自举数据扩充\n'
)

heading3('实验设计与评估方案')

body(
    '(a) 来自阿里巴巴运维日志的内部 AIOps 数据集（已与 SRE 团队初步确认数据访问意向，数据共享协议商定中），'
    '预估规模 ≥10,000 条故障事件，以 8:1:1 比例切分训练/验证/测试集。\n'
    '(b) 公开可用的 AIOps-2022 基准（保障可复现性）。\n'
    '覆盖三个子场景：数据库层故障、网络分区事件、容器资源耗尽；每条查询采样 k=8 条推理链。',
    bold_prefix='数据集\n'
)

body(
    '  ① 仅 LLM：通过 API 调用 72B，无本地 SLM。\n'
    '  ② SLM（无 CoT）：SLM 仅在 ⟨输入, 输出⟩ 对上微调，无推理链。\n'
    '  ③ SLM + 教师 CoT：SLM 在由 72B 作为教师生成的 CoT 上微调（oracle 上界）。\n'
    '  ④ 本方案：SLM 在答案验证 CoT 上微调，教师仅为任务可行的最小模型。',
    bold_prefix='对比基线\n'
)

body(
    '  ▸ 符号层：修复操作准确率（top-1 命中率）、推理延迟（p50/p99）、单查询计算成本（FLOPs / API token）。\n'
    '  ▸ 执行层：内存峰值恢复率、服务心跳通过率、滚动窗口错误率下降比、p99 延迟恢复比。\n'
    '  ▸ 统计显著性通过自助重采样（n=1000）评估。',
    bold_prefix='评估指标\n'
)

body(
    '云端 LLM（72B 或同等模型）充当 Planner，将故障事件分解为子任务并路由至对应的专家小智能体（SLM）；'
    '每个小智能体专为单一 AIOps 子场景训练（如数据库慢查询诊断、网络分区检测、容器 OOM 分诊），'
    '本地推理延迟 <100ms。\n'
    'Planner–专家小智能体分层设计同时保留了大模型的全局语义广度与小模型的窄域推理深度，'
    '端到端系统可在秒级内响应高频告警。',
    bold_prefix='系统架构：Planner + 专家小智能体\n'
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. 里程碑计划
# ══════════════════════════════════════════════════════════════════════════════
heading1('2.  里程碑计划')
heading2('2.1  项目周期')
body('一年，自《研究合作协议》生效之日起计。')

heading2('2.2  计划安排')

milestone_data = [
    ('M1', '流水线设计与数据准备。\n形式化三循环迭代算法；采集并预处理 AIOps ⟨输入, 输出⟩ 数据集；'
           '实现 CoT 候选生成与答案验证过滤器。',
     '数据集与过滤模块；技术报告。', '第 1–3 月'),
    ('M2', 'Cycle 2/3 训练与基准测试。\n（Cycle 2）在 D* 上 SFT+LoRA 训练得到 π₀；'
           '（Cycle 3）施加 AIOps-GRPO 强化推理质量，得到 π*；'
           '与通用 LLM 基线对比；专业化机制消融分析。',
     '训练好的 π₀ 和 π* 模型；消融实验报告；CCF-A 论文#1 草稿。', '第 4–7 月'),
    ('M3', 'Bootstrap 迭代与系统集成。\n对失败子任务运行 Bootstrap 回路；'
           '将各 π* 打包为专家小智能体；集成 Planner 与小智能体层；端到端延迟与准确率评估。',
     'Bootstrap 实验报告；开源专家小智能体套件；系统演示；CCF-A 论文#2 草稿。', '第 8–9 月'),
    ('M4', '评估、专利申请与最终交付。\n全系统评估；与阿里巴巴相关方开展用户研究；专利申请；项目结题报告。',
     '结题报告；已发表/投稿论文；专利申请（1–2 项）。', '第 10–12 月'),
]

body('表1  项目里程碑与交付物。', bold_prefix='')
tbl2 = doc.add_table(rows=1 + len(milestone_data), cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(tbl2, 14.4)
hdrs = ['阶段', '工作内容', '交付物', '时间段']
col_w2 = [Cm(1.4), Cm(6.0), Cm(5.0), Cm(2.0)]
for j, h in enumerate(hdrs):
    c = tbl2.rows[0].cells[j]
    c.width = col_w2[j]
    set_cell_margins(c)
    shade_cell(c, 'D4D4D4')
    p_h = c.paragraphs[0]
    r_h = p_h.add_run(h)
    r_h.font.bold = True; r_h.font.size = Pt(10); r_h.font.name = '黑体'
    r_h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_line_spacing(p_h, 1.15, 0)
for i, (m, work, deliv, period) in enumerate(milestone_data):
    row = tbl2.rows[i + 1]
    for j, (txt, w) in enumerate(zip([m, work, deliv, period], col_w2)):
        c = row.cells[j]
        c.width = w
        set_cell_margins(c)
        p_c = c.paragraphs[0]
        r_c = p_c.add_run(txt)
        r_c.font.size = Pt(10); r_c.font.name = '宋体'
        r_c._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if j == 0:
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_c.font.bold = True
        _apply_line_spacing(p_c, 1.15, 0)

# 甘特图（表格形式）
doc.add_paragraph()
body('图1  项目甘特图（四个里程碑）。', bold_prefix='')

gantt_rows = [
    ('任务', '1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12'),
    ('M1 流水线设计与数据准备',   '█', '█', '█', '', '', '', '', '', '', '', '', ''),
    ('M2 Cycle 2/3 训练与基准',  '', '', '', '█', '█', '█', '█', '', '', '', '', ''),
    ('M3 Bootstrap 与系统集成',  '', '', '', '', '', '', '', '█', '█', '', '', ''),
    ('M4 评估与最终交付',         '', '', '', '', '', '', '', '', '', '█', '█', '█'),
]
tbl_g = doc.add_table(rows=len(gantt_rows), cols=13)
tbl_g.style = 'Table Grid'
tbl_g.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(tbl_g, 14.4)
for i, row_data in enumerate(gantt_rows):
    for j, txt in enumerate(row_data):
        c = tbl_g.rows[i].cells[j]
        c.width = Cm(1.0) if j == 0 else Cm(1.0)
        if j == 0:
            c.width = Cm(5.4)
        else:
            c.width = Cm(0.75)
        set_cell_margins(c, top=40, start=60, bottom=40, end=60)
        para = c.paragraphs[0]
        r = para.add_run(txt)
        r.font.size = Pt(9)
        if i == 0:
            r.font.bold = True
            r.font.name = '黑体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            shade_cell(c, 'D4D4D4')
        else:
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if txt == '█':
                shade_cell(c, 'A0A0A0')
        if j > 0:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_line_spacing(para, 1.15, 0)

# ══════════════════════════════════════════════════════════════════════════════
#  3. 交付成果
# ══════════════════════════════════════════════════════════════════════════════
heading1('3.  交付成果')
body('本项目承诺如下最终交付成果：')

deliverables = [
    ('算法原型系统',
     '完整的三循环迭代 SLM 生产流水线：数据采集模块、CoT 候选生成与答案验证过滤器（Cycle 1）；'
     'SFT + LoRA 监督训练模块（Cycle 2）；AIOps-GRPO 策略优化模块，含奖励函数实现（Cycle 3）；Bootstrap 自举回路。'
     '源代码在 GitHub 开源发布，提供完整可复现实验环境。'),
    ('AIOps 专家小智能体套件',
     '针对 2–3 个 AIOps 子场景（数据库层故障、网络分区事件、容器资源耗尽）'
     '打包的专家小智能体原型，含 Planner 层集成文档与接入指南。'),
    ('设计文档',
     '提交 1 份总体设计文档（系统架构与流水线规格）及 1 份详细设计文档'
     '（算法、评估方案与部署指南），分别于 M1、M3 里程碑交付。'),
    ('论文目标',
     '投稿论文 2 篇至 CCF-A 或同等顶级会议（如 NeurIPS、ICML、ACL、ICSE、FSE），'
     '涵盖所提方法与 AIOps 评估结果，争取于项目结题前至少 1 篇完成录用。'),
    ('专利目标',
     '就冷启动 CoT 合成与专家小智能体生产提交发明专利申请 1–2 项（国内或国际）。'),
    ('技术指标',
     '所产出 SLM 修复操作准确率在 GPT-4 级 LLM 的 5% 以内，'
     '单次推理延迟降低 >10×（从 ~1s 降至 <100ms），单查询 token 成本降至零（本地推理）。'),
    ('业务目标',
     '系统在 2 个以上 AIOps 子场景完成原型部署与性能验证，'
     '并有文档记录相对仅 LLM 基线的延迟与准确率提升。'),
    ('项目验收报告', '项目结束时提交 1 份最终验收测试报告。'),
    ('学生联合培养与实习',
     '联合培养参与本项目的研究生 2–3 名；在项目期间提供赴阿里巴巴 SRE 团队实习机会'
     '（预计第 4–9 月，以阿里巴巴排期为准）。'),
]
for idx, (title, desc) in enumerate(deliverables, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _apply_line_spacing(p, 1.15, 4)
    r_t = p.add_run(f'{idx}. 【{title}】')
    r_t.font.bold = True; r_t.font.size = Pt(11); r_t.font.name = '宋体'
    r_t._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r_d = p.add_run(f'\n{desc}')
    r_d.font.size = Pt(11); r_d.font.name = '宋体'
    r_d._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
body('表2  交付成果汇总（含交付方式）。', bold_prefix='')
summary_data = [
    ('原型系统：三循环迭代冷启动 SLM 生产流水线', '源代码（开源）', '可复现；在 GitHub 发布'),
    ('AIOps 专家小智能体套件', '源代码 + 文档', '兼容 Planner 层接口'),
    ('总体设计文档与详细设计文档', '技术报告', '分别于 M1、M3 里程碑提交'),
    ('学术论文 2 篇', '学术论文', '目标 CCF-A 或同等顶级会议'),
    ('专利申请 1–2 项', '专利申请', '国内或国际专利'),
    ('中期与结题项目报告', '项目报告', '按约定时间表提交'),
    ('送研究生赴阿里巴巴 SRE 团队实习', '人员安排', '2–3 名博士生；预计第 4–9 月'),
]
tbl3 = doc.add_table(rows=1 + len(summary_data), cols=3)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(tbl3, 14.4)
col_w3 = [Cm(6.5), Cm(3.5), Cm(4.4)]
for j, h in enumerate(['说明', '交付方式', '备注']):
    c = tbl3.rows[0].cells[j]
    c.width = col_w3[j]
    set_cell_margins(c)
    shade_cell(c, 'D4D4D4')
    r_h = c.paragraphs[0].add_run(h)
    r_h.font.bold = True; r_h.font.size = Pt(10); r_h.font.name = '黑体'
    r_h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    _apply_line_spacing(c.paragraphs[0], 1.15, 0)
for i, row_data in enumerate(summary_data):
    for j, (txt, w) in enumerate(zip(row_data, col_w3)):
        c = tbl3.rows[i + 1].cells[j]
        c.width = w
        set_cell_margins(c)
        r_c = c.paragraphs[0].add_run(txt)
        r_c.font.size = Pt(10); r_c.font.name = '宋体'
        r_c._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        _apply_line_spacing(c.paragraphs[0], 1.15, 0)

# ══════════════════════════════════════════════════════════════════════════════
#  4. 人员、设备与经费预算
# ══════════════════════════════════════════════════════════════════════════════
heading1('4.  人员、设备与经费预算')
heading2('4.1  人员配置')

staff_data = [
    ('王茂林', '研究助理教授，\n香港城市大学', '12,000', '3 月', '0.25 FTE', '36,000', '负责人', '—'),
    ('<FILL: PhD 1>', '博士生，\n<FILL: Univ.>', '12,000', '12 月', '1.0 FTE', '144,000', '研究助理', '第 4–6 月'),
    ('<FILL: PhD 2>', '博士生，\n<FILL: Univ.>', '12,000', '12 月', '1.0 FTE', '144,000', '研究助理', '第 7–9 月'),
]
tbl4 = doc.add_table(rows=1 + len(staff_data) + 1, cols=8)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(tbl4, 14.4)
staff_hdrs = ['姓名', '背景', '月薪（元）', '周期', '折合全时', '合计', '角色', '实习']
for j, h in enumerate(staff_hdrs):
    c = tbl4.rows[0].cells[j]
    set_cell_margins(c, top=40, start=60, bottom=40, end=60)
    shade_cell(c, 'D4D4D4')
    r_h = c.paragraphs[0].add_run(h)
    r_h.font.bold = True; r_h.font.size = Pt(9); r_h.font.name = '黑体'
    r_h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_line_spacing(c.paragraphs[0], 1.15, 0)
for i, row_data in enumerate(staff_data):
    for j, txt in enumerate(row_data):
        c = tbl4.rows[i + 1].cells[j]
        set_cell_margins(c, top=40, start=60, bottom=40, end=60)
        r_c = c.paragraphs[0].add_run(txt)
        r_c.font.size = Pt(9); r_c.font.name = '宋体'
        r_c._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        _apply_line_spacing(c.paragraphs[0], 1.15, 0)
# 小计行
subtotal_row = tbl4.rows[len(staff_data) + 1]
c_label = subtotal_row.cells[0]
for k in range(1, 5):
    c_label.merge(subtotal_row.cells[k])
set_cell_margins(c_label, top=40, start=60, bottom=40, end=60)
r_l = c_label.paragraphs[0].add_run('人员费用小计（a）')
r_l.font.bold = True; r_l.font.size = Pt(9); r_l.font.name = '黑体'
r_l._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
c_label.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
_apply_line_spacing(c_label.paragraphs[0], 1.15, 0)
shade_cell(c_label, 'ECECEC')
c_val = subtotal_row.cells[5]
set_cell_margins(c_val, top=40, start=60, bottom=40, end=60)
shade_cell(c_val, 'ECECEC')
r_v = c_val.paragraphs[0].add_run('RMB 324,000')
r_v.font.bold = True; r_v.font.size = Pt(9); r_v.font.name = '黑体'
r_v._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
_apply_line_spacing(c_val.paragraphs[0], 1.15, 0)

body('注：实习时间段为预估值；具体日期以阿里巴巴 SRE 团队排期为准。')

heading2('4.2  设备与耗材')
equip_data = [
    ('GPU 云计算（A100 节点时长）', '—', '91,000', '用于 SLM 训练（第二至三阶段）'),
    ('本地推理服务器（租赁/购买）', '50,000', '50,000', 'MCP 端点部署与延迟测试'),
    ('学术会议差旅（2 场）', '15,000', '30,000', '国内及国际顶级会议参会'),
    ('开放获取出版费', '10,000', '20,000', 'ACL Anthology / NeurIPS 开放获取'),
]
tbl5 = doc.add_table(rows=1 + len(equip_data) + 1, cols=4)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(tbl5, 14.4)
col_w5 = [Cm(6.0), Cm(2.4), Cm(2.4), Cm(3.6)]
for j, h in enumerate(['项目', '单价（元）', '合计（元）', '备注']):
    c = tbl5.rows[0].cells[j]
    c.width = col_w5[j]
    set_cell_margins(c)
    shade_cell(c, 'D4D4D4')
    r_h = c.paragraphs[0].add_run(h)
    r_h.font.bold = True; r_h.font.size = Pt(10); r_h.font.name = '黑体'
    r_h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    _apply_line_spacing(c.paragraphs[0], 1.15, 0)
for i, row_data in enumerate(equip_data):
    for j, (txt, w) in enumerate(zip(row_data, col_w5)):
        c = tbl5.rows[i + 1].cells[j]
        c.width = w
        set_cell_margins(c)
        r_c = c.paragraphs[0].add_run(txt)
        r_c.font.size = Pt(10); r_c.font.name = '宋体'
        r_c._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        _apply_line_spacing(c.paragraphs[0], 1.15, 0)
subtotal_row2 = tbl5.rows[len(equip_data) + 1]
subtotal_row2.cells[0].merge(subtotal_row2.cells[1])
c_l2 = subtotal_row2.cells[0]
set_cell_margins(c_l2)
shade_cell(c_l2, 'ECECEC')
r_l2 = c_l2.paragraphs[0].add_run('科研经费小计（b）')
r_l2.font.bold = True; r_l2.font.size = Pt(10); r_l2.font.name = '黑体'
r_l2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
c_l2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
_apply_line_spacing(c_l2.paragraphs[0], 1.15, 0)
c_v2 = subtotal_row2.cells[2]
set_cell_margins(c_v2)
shade_cell(c_v2, 'ECECEC')
r_v2 = c_v2.paragraphs[0].add_run('191,000')
r_v2.font.bold = True; r_v2.font.size = Pt(10); r_v2.font.name = '黑体'
r_v2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
_apply_line_spacing(c_v2.paragraphs[0], 1.15, 0)

heading2('4.3  经费预算')
budget_data = [
    ('（a）人员费用合计', '详见第 4.1 节', '324,000'),
    ('（b）科研经费合计', '详见第 4.2 节', '191,000'),
    ('（c）间接费用', '学校管理费：16.5% × (a+b) = 16.5% × 515,000 ≈ 85,000', '85,000'),
]
tbl6 = doc.add_table(rows=1 + len(budget_data) + 1, cols=3)
tbl6.style = 'Table Grid'
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(tbl6, 14.4)
col_w6 = [Cm(4.0), Cm(7.4), Cm(3.0)]
for j, h in enumerate(['项目', '备注', '金额（元）']):
    c = tbl6.rows[0].cells[j]
    c.width = col_w6[j]
    set_cell_margins(c)
    shade_cell(c, 'D4D4D4')
    r_h = c.paragraphs[0].add_run(h)
    r_h.font.bold = True; r_h.font.size = Pt(10); r_h.font.name = '黑体'
    r_h._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    _apply_line_spacing(c.paragraphs[0], 1.15, 0)
for i, row_data in enumerate(budget_data):
    for j, (txt, w) in enumerate(zip(row_data, col_w6)):
        c = tbl6.rows[i + 1].cells[j]
        c.width = w
        set_cell_margins(c)
        r_c = c.paragraphs[0].add_run(txt)
        r_c.font.size = Pt(10); r_c.font.name = '宋体'
        r_c._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        _apply_line_spacing(c.paragraphs[0], 1.15, 0)
total_row = tbl6.rows[len(budget_data) + 1]
shade_cell(total_row.cells[0], 'ECECEC')
shade_cell(total_row.cells[1], 'ECECEC')
shade_cell(total_row.cells[2], 'ECECEC')
for j, (txt, w, bold) in enumerate(zip(
        ['合计', '（a）+（b）+（c）', '600,000'],
        col_w6, [True, False, True])):
    c = total_row.cells[j]
    c.width = w
    set_cell_margins(c)
    r = c.paragraphs[0].add_run(txt)
    r.font.bold = bold; r.font.size = Pt(10)
    r.font.name = '黑体' if bold else '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if bold else '宋体')
    _apply_line_spacing(c.paragraphs[0], 1.15, 0)

# ══════════════════════════════════════════════════════════════════════════════
#  5. 主要研究人员简介
# ══════════════════════════════════════════════════════════════════════════════
heading1('5.  主要研究人员简介')

p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 4)
r = p.add_run('王茂林（Maolin Wang）')
r.font.bold = True; r.font.size = Pt(12); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

body('研究助理教授，计算机科学系，香港城市大学（City University of Hong Kong）\n'
     '电子邮件：maolwang2@cityu.edu.hk\n'
     '个人主页：https://scholars.cityu.edu.hk/en/persons/maolwang2')

p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 3)
r = p.add_run('教育背景：')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
bullet('博士，数据科学方向，香港城市大学（2021–2025）\n（导师：赵翔宇 教授；共同导师：Ruocheng Guo 博士、王俊晖 教授）')
bullet('硕士，计算机科学系，电子科技大学（2018–2021）（导师：徐增林 教授）')
bullet('学士，英才实验学院，电子科技大学（2014–2018）')

p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 3)
r = p.add_run('研究方向：')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
body('大语言模型、知识蒸馏、AIOps 与智能运维、自动机器学习、智能软件系统。')

p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 3)
r = p.add_run('代表性成果（近五年）：')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

pubs = [
    'Wang M. et al. "FLUID-MMRec: Stein-Guided Entropic Flow for Multi-Modal Sequential Recommendation." KDD 2025（CCF-A）。[第一作者]',
    'Wang M. et al. "DANCE: Resource-Efficient Neural Architecture Search with Data-Aware and Continuous Adaptation." KDD 2025 & IJCAI 2025（CCF-A）。[第一作者]',
    'Peng J., Wang M. et al. "Stepwise Reasoning Error Disruption Attack of LLMs." ACL 2025（CCF-A）。',
    'Liu Z., Liu Q., Wang M. et al. "SIGMA: Selective Gated Mamba for Sequential Recommendation." AAAI 2025（CCF-A）。',
    'Han X., Wang M. et al. "Data Efficient Adaptation in Large Language Models via Continuous Low-Rank Fine-Tuning." NeurIPS 2025（CCF-A）。',
]
for i, pub in enumerate(pubs, 1):
    numbered(pub, i)

p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 3)
r = p.add_run('代表性项目与获奖：')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
bullet('KDD 2025 最佳论文亚军奖（Best Paper Award Runner-Up），ACM SIGKDD，2025年。')
bullet('ACM SIGKDD 学生旅行奖（KDD 2025 Student Travel Award），2025年。')
bullet('代表性产业落地：模型压缩成果（EI-BERT，1.9MB）已在支付宝多条业务线上线，成为基础能力。')

p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 3)
r = p.add_run('学术服务：')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
bullet('程序委员会委员 / 审稿人：NeurIPS、KDD、WWW、AAAI、IJCAI、ACL、SIGIR、RecSys。')

p = doc.add_paragraph()
_apply_line_spacing(p, 1.15, 3)
r = p.add_run('与本申请高度相关的具体经历：')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '黑体'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
numbered('<FILL: 在该方向于[列举 CCF-A 会议及年份]发表论文 N 篇，其中 N 篇以第一作者或通讯作者身份发表。代表论文：[1] … [2] …>', 1)
numbered('<FILL: 在[公司]研究合作/实习期间（[时长]），作为核心成员参与[项目名称]，主导[具体模型/组件]在[N]条业务线（[产品名称]）上的部署，取得[具体量化成果]。>', 2)
numbered('<FILL: 在 CCF-A 会议（[列举会议及年份]）就[相关主题]发表[教程/特邀报告]。就该主题撰写综述论文 [题目, arXiv/会议, 年份]。>', 3)

# ══════════════════════════════════════════════════════════════════════════════
#  6. 参考文献（与 PDF 保持一致）
# ══════════════════════════════════════════════════════════════════════════════
heading1('6.  参考文献')

refs = [
    '[1] Huang R. et al. "AIOps Challenge 2022: Cloud Microservice Anomaly Detection." 2022.',
    '[2] Gu Y. et al. "MiniLLM: Knowledge Distillation of Large Language Models." ICLR 2024.',
    '[3] Agarwal R. et al. "On-Policy Distillation of Language Models." NeurIPS 2024.',
    '[4] Zelikman E. et al. "STaR: Bootstrapping Reasoning With Reasoning." NeurIPS 2022.',
    '[5] Wei J. et al. "Chain-of-Thought Prompting Elicits Reasoning in Large Language Models." NeurIPS 2022.',
    '[6] Hinton G. et al. "Distilling the Knowledge in a Neural Network." NIPS Workshop 2014.',
    '[7] Chen L. et al. "Logos: Empowering Molecule Design with Expert-Guided Reasoning." 2024.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    _apply_line_spacing(p, 1.15, 3)
    r = p.add_run(ref)
    r.font.size = Pt(10); r.font.name = '宋体'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 保存 ──────────────────────────────────────────────────────────────────────
out = '/home/user/LatexAliBaBa/申请书_阿里巴巴AI研究资助计划2026.docx'
doc.save(out)
print(f'已生成：{out}')
